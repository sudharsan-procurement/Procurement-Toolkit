"""
Stage 1 — EXTRACT.

Goal: turn an uploaded file (digital PDF or .docx) into plain text, line by line.

"Digital" means the text is really stored in the file (you can select it in a PDF
reader). Scanned PDFs are just photos of text and need OCR — that is deliberately
NOT handled here yet; it is a later stage in the roadmap. If someone uploads a
scanned PDF, we detect that we got almost no text back and tell them clearly.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field


@dataclass
class ExtractResult:
    text: str            # the full document text
    kind: str            # "pdf" or "docx"
    looks_scanned: bool  # True if we suspect this is a scanned image PDF
    note: str            # human-readable status message
    # Per-page OCR confidence, populated only when OCR ran, e.g.
    # [{"page": 1, "confidence": 98.0}, ...]. Empty for non-OCR files.
    ocr_pages: list = field(default_factory=list)


def extract(file_bytes: bytes, filename: str) -> ExtractResult:
    """Dispatch on file extension and return the extracted text."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if name.endswith(".docx"):
        return _extract_docx(file_bytes)
    if name.endswith(".xlsx"):
        return _extract_xlsx(file_bytes)
    if name.endswith(".csv"):
        return _extract_csv(file_bytes)
    if name.endswith((".txt", ".md", ".markdown")):
        return _extract_text(file_bytes, name)
    if name.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")):
        return _extract_image(file_bytes)
    if name.endswith(".doc"):
        raise ValueError(
            "Old-style .doc files aren't supported. Please save as .docx or PDF."
        )
    if name.endswith(".xls"):
        raise ValueError(
            "Old-style .xls files aren't supported. Please save as .xlsx."
        )
    raise ValueError(
        f"Unsupported file type: {filename!r}. Upload a PDF, Word (.docx), "
        "Excel (.xlsx), CSV (.csv), text (.txt), Markdown (.md), or an image "
        "(.png/.jpg/.tiff)."
    )


def _extract_pdf(file_bytes: bytes) -> ExtractResult:
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # extract_text returns None for pages with no embedded text (e.g. scans)
            pages_text.append(page.extract_text() or "")

    text = "\n".join(pages_text).strip()

    # Heuristic: a digital contract has plenty of characters per page. If we got
    # almost nothing back, it is almost certainly a scan that needs OCR.
    chars_per_page = len(text) / max(len(pages_text), 1)
    looks_scanned = chars_per_page < 50

    if not looks_scanned:
        note = f"Read {len(pages_text)} page(s) of digital PDF text."
        return ExtractResult(text=text, kind="pdf", looks_scanned=False, note=note)

    # --- Scanned PDF: try OCR (reading text out of the page images). ---
    from .ocr import tesseract_available, ocr_pdf_with_confidence

    if tesseract_available():
        ocr_text, ocr_pages = ocr_pdf_with_confidence(file_bytes)
        if len(ocr_text) > len(text):
            note = f"Scanned PDF read with OCR ({len(ocr_pages)} page(s))."
            return ExtractResult(text=ocr_text, kind="pdf", looks_scanned=False,
                                 note=note, ocr_pages=ocr_pages)

    # OCR unavailable, or it found nothing useful.
    note = (
        "This PDF appears to be SCANNED and OCR couldn't read it (the scan may be "
        "low quality, or the OCR engine isn't available). Try a clearer scan or a "
        "digital PDF."
    )
    return ExtractResult(text=text, kind="pdf", looks_scanned=True, note=note)


def _extract_docx(file_bytes: bytes) -> ExtractResult:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(file_bytes))

    lines: list[str] = []
    for para in document.paragraphs:
        lines.append(para.text)

    # Also pull table cells out as lines so numbers in tables aren't lost.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    note = "Read Word (.docx) text" + (
        " (including simple tables)." if document.tables else "."
    )
    return ExtractResult(text=text, kind="docx", looks_scanned=False, note=note)


def _extract_xlsx(file_bytes: bytes) -> ExtractResult:
    import openpyxl

    # data_only=True returns calculated VALUES (not formulas) so numbers compare
    # correctly. read_only=True keeps memory low on the free host.
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)

    lines: list[str] = []
    for ws in wb.worksheets:
        if len(wb.worksheets) > 1:
            lines.append(f"[Sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            # Each row becomes one line: "cell | cell | cell" (blanks dropped).
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    note = f"Read Excel (.xlsx): {len(wb.worksheets)} sheet(s), one row per line."
    return ExtractResult(text=text, kind="xlsx", looks_scanned=False, note=note)


def _extract_csv(file_bytes: bytes) -> ExtractResult:
    import csv

    # utf-8-sig also strips a leading byte-order mark that Excel-exported CSVs
    # often add; Latin-1 is a last-resort fallback so nothing ever crashes.
    try:
        text_data = file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        text_data = file_bytes.decode("latin-1", errors="replace")

    lines: list[str] = []
    for row in csv.reader(io.StringIO(text_data)):
        # Each row becomes one line: "cell | cell | cell" (blanks dropped), so
        # numbers in the cells are still picked up and compared.
        cells = [c.strip() for c in row if c is not None and c.strip()]
        if cells:
            lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    note = f"Read CSV: {len(lines)} row(s), one row per line."
    return ExtractResult(text=text, kind="csv", looks_scanned=False, note=note)


def _extract_text(file_bytes: bytes, name: str) -> ExtractResult:
    # Plain text and Markdown are already text — just decode the bytes. We try
    # UTF-8 first (the modern default) and fall back to Latin-1 so older files
    # never crash the app.
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")

    kind = "md" if name.endswith((".md", ".markdown")) else "txt"
    label = "Markdown" if kind == "md" else "plain text"
    return ExtractResult(text=text.strip(), kind=kind, looks_scanned=False,
                         note=f"Read {label} file.")


def _extract_image(file_bytes: bytes) -> ExtractResult:
    # An image is just a picture, so the only way to get text out is OCR.
    from .ocr import tesseract_available, ocr_image_with_confidence

    if not tesseract_available():
        raise ValueError(
            "Image files need OCR to read, but the OCR engine isn't available "
            "on this server."
        )

    text, ocr_pages = ocr_image_with_confidence(file_bytes)
    note = "Read image text with OCR."
    if not text:
        note = "OCR found no readable text in this image (it may be blank or unclear)."
    return ExtractResult(text=text, kind="image", looks_scanned=False, note=note,
                         ocr_pages=ocr_pages)
