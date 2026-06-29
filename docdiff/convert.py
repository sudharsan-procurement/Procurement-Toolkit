"""
Converters — the extra tools that live alongside the comparison engine.

    pdf_to_word            : turn a digital PDF into an editable Word (.docx)
    word_tables_to_excel   : pull every table out of a Word file into Excel (.xlsx)

Both run locally and free. Like the rest of the app, the PDF converter works on
DIGITAL PDFs (real selectable text); scanned/photographed PDFs need OCR, which is
a later roadmap item.
"""

from __future__ import annotations

import io
import os
import tempfile


def pdf_to_word(pdf_bytes: bytes) -> bytes:
    """Convert a digital PDF into a Word .docx and return the new file's bytes."""
    from pdf2docx import Converter

    # pdf2docx works with files on disk, so we use a short-lived temp folder that
    # is deleted automatically when we're done.
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, "input.pdf")
        docx_path = os.path.join(tmp, "output.docx")

        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)

        converter = Converter(pdf_path)
        try:
            converter.convert(docx_path)  # convert all pages
        finally:
            converter.close()

        with open(docx_path, "rb") as f:
            return f.read()


def word_tables_to_excel(
    docx_bytes: bytes, separate_sheets: bool = True
) -> tuple[bytes, int]:
    """
    Extract every table from a Word .docx into an Excel workbook.

    `separate_sheets` controls the layout:
        True  -> each Word table gets its own sheet ("Table 1", "Table 2", …)
        False -> all tables stacked on ONE sheet, each under a "Table N" label
                 with a blank spacer row between them.

    Returns the Excel file's bytes plus the number of tables found (0 if none).
    """
    import docx  # python-docx
    from openpyxl import Workbook

    document = docx.Document(io.BytesIO(docx_bytes))
    tables = document.tables
    count = len(tables)

    wb = Workbook()
    wb.remove(wb.active)  # drop the default blank sheet; we add our own

    if count == 0:
        ws = wb.create_sheet(title="No tables found")
        ws.append(["No tables were found in this Word document."])

    elif separate_sheets:
        # One sheet per table. Sheet names stay well under Excel's 31-char limit.
        for i, table in enumerate(tables, start=1):
            ws = wb.create_sheet(title=f"Table {i}")
            for row in table.rows:
                ws.append([cell.text.strip() for cell in row.cells])

    else:
        # All tables on a single sheet, each labelled and separated by a blank row.
        ws = wb.create_sheet(title="All tables")
        for i, table in enumerate(tables, start=1):
            ws.append([f"Table {i}"])
            for row in table.rows:
                ws.append([cell.text.strip() for cell in row.cells])
            ws.append([])  # blank spacer row between tables

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue(), count


def extract_word_tables(docx_bytes: bytes) -> list[list[list[str]]]:
    """
    Read every table from a Word .docx into a plain nested list:
        [ table ][ row ][ cell-text ]

    This is the "preview" half of Word -> Excel: the app shows these tables and
    lets the user pick which to keep, rename, and merge, then calls
    build_tables_excel() with the chosen specs.
    """
    import docx  # python-docx

    document = docx.Document(io.BytesIO(docx_bytes))
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)
    return tables


def _safe_sheet_name(name: str, used: set[str]) -> str:
    """Return an Excel-safe, unique sheet name (<=31 chars, no []:*?/\\ )."""
    for ch in "[]:*?/\\":
        name = name.replace(ch, "-")
    name = name.strip()[:31] or "Sheet"
    candidate = name
    n = 2
    while candidate.lower() in used:
        suffix = f"_{n}"
        candidate = name[:31 - len(suffix)] + suffix
        n += 1
    used.add(candidate.lower())
    return candidate


def build_tables_excel(specs: list[dict], separate_sheets: bool = True) -> bytes:
    """
    Build an Excel workbook from already-prepared table specs.

    `specs` is a list of dicts: {"name": <sheet/label>, "rows": <list of rows>}.
    The caller (the app) has already applied the user's keep / rename / merge
    choices, so this just writes them out, honouring the two layout modes:
        separate_sheets=True  -> one sheet per spec (named by spec["name"])
        separate_sheets=False -> all specs stacked on one sheet, each under its label
    """
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)

    if not specs:
        ws = wb.create_sheet(title="No tables")
        ws.append(["No tables were selected for export."])

    elif separate_sheets:
        used: set[str] = set()
        for i, spec in enumerate(specs, start=1):
            title = _safe_sheet_name(spec.get("name") or f"Table_{i}", used)
            ws = wb.create_sheet(title=title)
            for row in spec["rows"]:
                ws.append(row)

    else:
        ws = wb.create_sheet(title="All tables")
        for i, spec in enumerate(specs, start=1):
            ws.append([spec.get("name") or f"Table {i}"])
            for row in spec["rows"]:
                ws.append(row)
            ws.append([])  # blank spacer row between tables

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
